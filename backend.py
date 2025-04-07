# backend.py (Modified for DOCX Input, Translation, and Merging - Updated with Merge Fix)

import io
import google.generativeai as genai
from docx import Document # For reading/writing .docx
# --- NEW: Import docxcompose for merging ---
from docxcompose.composer import Composer
# ---
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import logging
import os
# import streamlit as st # No longer needed in backend if Vision credentials removed
# import json # No longer needed if Vision credentials removed
# from google.cloud import vision # REMOVED

# --- Configure Logging ---
# Sets up basic logging to track the script's execution and potential errors.
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(module)s - %(message)s')

# --- REMOVED: Runtime Credentials Setup for Vision API ---
# No longer needed as this version focuses on Gemini API Key authentication.

# --- NEW: Function to extract text from DOCX, skipping headers/footers ---
def extract_text_from_docx(docx_file_obj):
    """
    Extracts text from the main body of a DOCX file object,
    attempting to skip headers and footers.

    Args:
        docx_file_obj: A file-like object (e.g., BytesIO, UploadedFile) representing the DOCX file.

    Returns:
        str: The extracted text, joined by newlines.
             Returns an error string starting with "Error:" if a failure occurs.
    """
    try:
        # Ensure the file pointer is at the beginning before reading
        docx_file_obj.seek(0)
        # Load the document using python-docx
        document = Document(docx_file_obj)
        full_text = []

        # Iterate through paragraphs in the main document body.
        # This approach implicitly skips headers/footers which are stored in different parts of the DOCX structure.
        for para in document.paragraphs:
            full_text.append(para.text)

        # Note: This simple approach doesn't explicitly handle footnotes,
        # text boxes, or complex table content. Further refinement might be needed
        # for documents with very complex layouts.
        extracted_content = "\n".join(full_text)
        logging.info(f"Extracted {len(extracted_content)} characters from DOCX main body.")
        return extracted_content

    except Exception as e:
        # Log the error and return an informative message to the frontend.
        logging.error(f"Error extracting text from DOCX: {e}", exc_info=True)
        # Provide a more informative error message if possible
        return f"Error: Failed to process DOCX file. Check if it's a valid .docx format. Details: {e}"


# --- Gemini Processing (Simplified Prompt Focus - Translation Only) ---
def process_text_with_gemini(api_key: str, raw_text: str, rules_prompt: str, model_name: str):
    """
    Sends text (extracted from DOCX) to Google Gemini for translation based on provided rules.

    Args:
        api_key (str): The Google Gemini API key.
        raw_text (str): The text extracted from the DOCX (expected to be Urdu, Farsi, or English).
        rules_prompt (str): User-defined instructions for the translation (e.g., translate to Arabic).
        model_name (str): The specific Gemini model ID (e.g., 'gemini-1.5-flash-latest').

    Returns:
        str: The translated Arabic text from Gemini, an empty string if input was empty or translation failed gracefully,
             or an error string starting with "Error:" if a critical failure occurs.
    """
    # Basic input validation
    if not api_key:
        logging.error("Gemini API key is missing.")
        return "Error: Gemini API key is missing."

    if not raw_text or not raw_text.strip():
        # If there's no text to translate (e.g., empty DOCX), don't call the API.
        logging.warning("Skipping Gemini call: No text extracted from DOCX.")
        return "" # Return empty string, indicating nothing was translated.

    if not model_name:
        logging.error("Gemini model name is missing.")
        return "Error: Gemini model name not specified."

    try:
        # Configure the Gemini client library with the API key.
        genai.configure(api_key=api_key)
        logging.info(f"Initializing Gemini model: {model_name}")
        # Instantiate the generative model.
        model = genai.GenerativeModel(model_name)

        # Construct the full prompt for the Gemini model, combining the rules and the text.
        full_prompt = f"""
        **Instructions:**
        {rules_prompt}

        **Text to Process:**
        ---
        {raw_text}
        ---

        **Output:**
        Return ONLY the processed text (Arabic translation) according to the instructions.
        """

        logging.info(f"Sending request to Gemini model: {model_name} for translation. Text length: {len(raw_text)}")
        # Send the request to the Gemini API.
        response = model.generate_content(full_prompt)

        # --- Handle potential issues with the response ---
        # Check if the response contains any parts (content).
        if not response.parts:
            block_reason = None
            safety_ratings = None
            # Check for safety feedback (blocking reasons).
            if hasattr(response, 'prompt_feedback'):
                block_reason = getattr(response.prompt_feedback, 'block_reason', None)
                safety_ratings = getattr(response.prompt_feedback, 'safety_ratings', None)

            if block_reason:
                # If content was blocked by safety filters.
                block_reason_msg = f"Content blocked by Gemini safety filters. Reason: {block_reason}"
                logging.error(f"Gemini request ({model_name}) blocked. Reason: {block_reason}. Ratings: {safety_ratings}")
                return f"Error: {block_reason_msg}"
            else:
                # If the response is empty but not explicitly blocked (e.g., model finished unexpectedly).
                finish_reason_obj = getattr(response, 'prompt_feedback', None)
                finish_reason = getattr(finish_reason_obj, 'finish_reason', 'UNKNOWN') if finish_reason_obj else 'UNKNOWN'
                logging.warning(f"Gemini ({model_name}) returned no parts (empty response). Finish Reason: {finish_reason}")
                return "" # Return empty string as no translation was provided.

        # Extract the translated text from the response.
        processed_text = response.text
        logging.info(f"Successfully received translation from Gemini ({model_name}). Length: {len(processed_text)}")
        return processed_text

    except Exception as e:
        # Catch any other exceptions during API interaction.
        logging.error(f"Error interacting with Gemini API ({model_name}): {e}", exc_info=True)
        return f"Error: Failed to process text with Gemini ({model_name}). Details: {e}"


# --- NEW: Create SINGLE Word Document with Arabic Text ---
def create_arabic_word_doc_from_text(arabic_text: str, filename: str):
    """
    Creates a single Word document (.docx) in memory containing the translated Arabic text.
    Sets paragraph alignment to right and text direction to RTL using Arial font.

    Args:
        arabic_text (str): The translated Arabic text received from Gemini.
        filename (str): Original filename, used for context in placeholders/logs if translation is empty.

    Returns:
        io.BytesIO: A BytesIO stream containing the Word document data, ready for merging or download.
                    Returns None if a critical error occurs during document creation.
    """
    try:
        # Create a new blank Word document.
        document = Document()
        # --- Set default styles for Arabic (RTL, Font) ---
        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial' # Set default font to Arial
        font.rtl = True     # Enable Right-to-Left for the default font

        # --- Set complex script font using low-level oxml ---
        # This ensures Arabic characters render correctly with the chosen font.
        style_element = style.element
        # Find or create the run properties element (<w:rPr>)
        rpr_elements = style_element.xpath('.//w:rPr')
        rpr = rpr_elements[0] if rpr_elements else OxmlElement('w:rPr')
        if not rpr_elements: style_element.append(rpr)
        # Find or create the font definition element (<w:rFonts>)
        font_name_element = rpr.find(qn('w:rFonts'))
        if font_name_element is None:
            font_name_element = OxmlElement('w:rFonts')
            rpr.append(font_name_element)
        # Set the complex script ('cs') font attribute to Arial
        font_name_element.set(qn('w:cs'), 'Arial')

        # --- Set default paragraph format ---
        paragraph_format = style.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT # Right-align paragraphs
        paragraph_format.right_to_left = True                 # Set paragraph direction to RTL

        # --- Add the translated Arabic text ---
        if arabic_text and arabic_text.strip():
            # Split the translated text into lines and add each as a paragraph.
            lines = arabic_text.strip().split('\n')
            for line in lines:
                if line.strip(): # Avoid adding empty paragraphs
                    paragraph = document.add_paragraph(line.strip())
                    # Explicitly set format for the added paragraph (redundant with style but safe)
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph.paragraph_format.right_to_left = True
                    # Ensure runs within the paragraph also use RTL/Complex Script settings
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.rtl = True
                        run.font.complex_script = True # Important for correct rendering
        else:
            # --- Handle cases where translation is empty or failed ---
            # Add a placeholder message to indicate which file this section corresponds to.
            empty_msg = f"[No translation generated or text extracted for '{os.path.basename(filename)}']"
            paragraph = document.add_paragraph(empty_msg)
            paragraph.italic = True # Italicize the placeholder message
            # Apply RTL formatting even to the placeholder for consistency
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.paragraph_format.right_to_left = True
            for run in paragraph.runs:
                run.font.name = 'Arial' # Use standard font for placeholder
                run.font.rtl = False    # LTR for the English placeholder text
                run.font.complex_script = False

        # --- Save document to an in-memory BytesIO stream ---
        doc_stream = io.BytesIO()
        document.save(doc_stream)
        doc_stream.seek(0) # Reset stream position to the beginning for reading later
        logging.info(f"Successfully created intermediate Word doc stream for '{filename}'.")
        return doc_stream

    except Exception as e:
        logging.error(f"Error creating intermediate Word document for '{filename}': {e}", exc_info=True)
        return None # Indicate failure


# --- RE-INTRODUCED: Merging Function using docxcompose (Corrected) ---
def merge_word_documents(doc_streams_data: list[tuple[str, io.BytesIO]]):
    """
    Merges multiple Word documents (provided as BytesIO streams) into one single document,
    adding a page break before appending each subsequent document.

    Args:
        doc_streams_data: A list of tuples, where each tuple contains:
                          (original_filename_str, word_doc_bytes_io_obj).
                          The list should be in the desired merge order.

    Returns:
        io.BytesIO: A BytesIO stream containing the final merged Word document data.
                    Returns None if no documents are provided or if a merging error occurs.
    """
    if not doc_streams_data:
        logging.warning("No document streams provided for merging.")
        return None

    try:
        # --- Initialize Composer with the first document ---
        first_filename, first_stream = doc_streams_data[0]
        first_stream.seek(0) # Ensure stream is at the start
        # Load the first document using python-docx. This becomes the base/master document.
        master_doc = Document(first_stream)
        # Create a Composer instance, using the first document as the base.
        composer = Composer(master_doc)
        logging.info(f"Initialized merger with base document from '{first_filename}'.")

        # --- Append remaining documents ---
        if len(doc_streams_data) > 1:
            # Loop through the rest of the document streams (starting from the second one).
            for i in range(1, len(doc_streams_data)):
                filename, stream = doc_streams_data[i]
                stream.seek(0) # Ensure stream is at the start
                logging.info(f"Preparing to append content from '{filename}'...")
                # Load the subsequent document that needs to be appended.
                sub_doc = Document(stream)

                # --- FIX: Add page break directly to the master_doc object ---
                # The page break is added to the end of the current master_doc content
                # *before* the content of sub_doc is appended by the composer.
                master_doc.add_page_break()
                # --- END FIX ---

                logging.info(f"Appending content from '{filename}'...")
                # Use the composer's append method to add the content of sub_doc to master_doc.
                composer.append(sub_doc)
                logging.info(f"Successfully appended content from '{filename}'.")

        # --- Save the final merged document ---
        merged_stream = io.BytesIO()
        # Use the composer's save method, which saves the modified master_doc.
        composer.save(merged_stream)
        merged_stream.seek(0) # Reset stream position for reading
        logging.info(f"Successfully merged {len(doc_streams_data)} documents.")
        return merged_stream

    except Exception as e:
        # Log any errors that occur during the merging process.
        logging.error(f"Error merging Word documents using docxcompose: {e}", exc_info=True)
        return None # Indicate failure

# --- REMOVED: append_text_to_document function ---
# This function is not needed for the merge strategy using docxcompose.
